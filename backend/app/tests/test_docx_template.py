from pathlib import Path

import pytest
from docx import Document

from app.docx_template import render_docx_template


def _create_template(tmp_path: Path, content: str) -> Path:
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph(content)
    doc.save(template_path)
    return template_path


def test_render_docx_template_with_mapping(tmp_path: Path) -> None:
    template_path = _create_template(tmp_path, "Bonjour {{ user.name }} !")

    output_path = render_docx_template(template_path, {"user": {"name": "Alice"}})

    rendered = Document(output_path)
    assert rendered.paragraphs[0].text == "Bonjour Alice !"
    assert output_path.name == "template_filled.docx"


def test_render_docx_template_from_json_string(tmp_path: Path) -> None:
    template_path = _create_template(tmp_path, "{{ title }} - {{ meta.author }}")
    json_data = '{"title": "Rapport", "meta": {"author": "Bob"}}'

    output_path = render_docx_template(template_path, json_data, tmp_path / "rapport.docx")

    rendered = Document(output_path)
    assert rendered.paragraphs[0].text == "Rapport - Bob"


def test_render_docx_template_invalid_input(tmp_path: Path) -> None:
    template_path = _create_template(tmp_path, "{{ field }}")

    with pytest.raises(TypeError):
        render_docx_template(template_path, ["not", "a", "mapping"])


def test_render_docx_template_missing_file(tmp_path: Path) -> None:
    missing_template = tmp_path / "missing.docx"

    with pytest.raises(FileNotFoundError):
        render_docx_template(missing_template, {"field": "value"})
