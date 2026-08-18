from __future__ import annotations

import datetime
import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(23, 63, 56)
ACCENT = RGBColor(8, 126, 112)
MUTED = RGBColor(100, 116, 139)


def _font(run, *, size: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size, run.bold = Pt(size), bold
    if color:
        run.font.color.rgb = color


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.alignment, table.autofit = WD_TABLE_ALIGNMENT.LEFT, False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_properties.append(cant_split)
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)


def _answer_text(value: Any) -> str:
    if value in (None, ""):
        return ""
    if value is True:
        return "Oui"
    if value is False:
        return "Non"
    return str(value)


def _plain_markdown(value: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(`+|\*\*|__|\*|_)", "", value)
    return value.strip()


def _add_markdown(document: Document, content: str, *, document_title: str) -> None:
    """Render the course Markdown structures used by laboratory sources."""
    lines = content.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index].strip()
        if not raw:
            index += 1
            continue
        if raw.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [_plain_markdown(cell) for cell in lines[index].strip().strip("|").split("|")]
                rows.append(cells)
                index += 1
            if len(rows) >= 2:
                data_rows = [rows[0], *rows[2:]]
                column_count = max(len(row) for row in data_rows)
                table = document.add_table(rows=len(data_rows), cols=column_count)
                table.style = "Table Grid"
                for row_index, source_row in enumerate(data_rows):
                    for column_index in range(column_count):
                        table.cell(row_index, column_index).text = source_row[column_index] if column_index < len(source_row) else ""
                        if row_index == 0:
                            _shade(table.cell(row_index, column_index), "E8F1EF")
                            for run in table.cell(row_index, column_index).paragraphs[0].runs:
                                _font(run, size=9.5, bold=True, color=INK)
                _set_table_geometry(table, [9360 // column_count] * column_count)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", raw)
        if heading:
            text = _plain_markdown(heading.group(2))
            if not (len(heading.group(1)) == 1 and text == document_title):
                document.add_heading(text, level=min(len(heading.group(1)), 3))
            index += 1
            continue
        if re.fullmatch(r"-{3,}", raw):
            index += 1
            continue
        if re.match(r"^[-*+]\s+", raw):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(_plain_markdown(re.sub(r"^[-*+]\s+", "", raw)))
            index += 1
            continue
        paragraph_lines = [raw]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "|", "- ", "* ", "+ ")):
                break
            paragraph_lines.append(candidate)
            index += 1
        document.add_paragraph(_plain_markdown(" ".join(paragraph_lines)))


def build_lab_attempt_docx(*, title: str, student_name: str, definition: dict[str, Any],
                           answers: dict[str, Any], validations: dict[str, Any],
                           status: str, updated_at: datetime.datetime | None) -> bytes:
    """Create a compact-reference-guide snapshot of the current attempt."""
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.25
    for style_name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 12, 10, 5)):
        style = document.styles[style_name]
        style.font.name, style.font.size, style.font.color.rgb = "Calibri", Pt(size), ACCENT
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    _font(kicker.add_run("COPIE ÉTUDIANTE · LABORATOIRE"), size=9, bold=True, color=ACCENT)
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    _font(heading.add_run(title), size=24, bold=True, color=INK)
    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(16)
    date_label = updated_at.astimezone().strftime("%Y-%m-%d %H:%M") if updated_at else "—"
    state_label = {"in_progress": "En cours", "submitted": "Remise", "evaluated": "Évaluée"}.get(status, status)
    _font(metadata.add_run(f"Étudiant : {student_name}\nÉtat : {state_label} · Dernier enregistrement : {date_label}"), size=10, color=MUTED)

    for block in definition.get("blocks", []):
        if block.get("type") == "markdown":
            _add_markdown(document, str(block.get("content", "")), document_title=title)
            continue
        field = block["field"]
        label = str(field.get("label", field.get("id", "Réponse")))
        if field.get("type") == "teacher_validation":
            document.add_heading(label, level=2)
            validation = validations.get(field["id"], {})
            text = "Validée" if validation.get("approved") else "En attente de validation"
            if validation.get("teacher_name"):
                text += f" par {validation['teacher_name']}"
            document.add_paragraph(text)
            continue
        value = answers.get(field["id"])
        if field.get("type") in {"table", "matrix"}:
            document.add_heading(label, level=2)
            visible = set(field.get("visible_columns", []))
            columns = [column for column in field.get("columns", []) if not visible or column["id"] in visible]
            table = document.add_table(rows=1, cols=len(columns) + 1)
            table.style = "Table Grid"
            header = table.rows[0].cells
            header[0].text = ""
            for index, column in enumerate(columns, 1):
                header[index].text = column.get("label", column["id"])
            for cell in header:
                _shade(cell, "E8F1EF")
                for run in cell.paragraphs[0].runs:
                    _font(run, size=9.5, bold=True, color=INK)
            cells = value if isinstance(value, dict) else {}
            for row in field.get("rows", []):
                word_row = table.add_row().cells
                word_row[0].text = row.get("label", row["id"])
                for index, column in enumerate(columns, 1):
                    answer = _answer_text(cells.get(f"{row['id']}.{column['id']}"))
                    unit = column.get("unit")
                    word_row[index].text = f"{answer} {unit}" if unit and answer else answer
            first, remaining = 2600, 6760
            _set_table_geometry(table, [first] + [remaining // max(len(columns), 1)] * len(columns))
        else:
            document.add_heading(label, level=2)
            answer = _answer_text(value)
            if field.get("unit") and answer:
                answer = f"{answer} {field['unit']}"
            paragraph = document.add_paragraph(answer)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(footer.add_run("Copie générée depuis edxo"), size=8.5, color=MUTED)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
